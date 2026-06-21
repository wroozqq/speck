import docx
from docx.shared import Inches, Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_borders_text(section):
    sectPr = section._sectPr
    # Remove existing pgBorders if any
    for el in list(sectPr):
        if el.tag.endswith('pgBorders'):
            sectPr.remove(el)
            
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offset-from'), 'text')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8') # 1 pt
        border.set(qn('w:space'), '10') # 10 pt space
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)

doc = docx.Document()
section = doc.sections[0]
section.top_margin = Cm(1.0)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(0.5)

add_page_borders_text(section)

footer = section.footer
# Create a table in footer
table = footer.add_table(rows=1, cols=3, width=Cm(18.5))
table.columns[0].width = Cm(5.0)
table.columns[1].width = Cm(8.5)
table.columns[2].width = Cm(5.0)

row = table.rows[0]
row.cells[0].text = "Left Cell"
row.cells[1].text = "Middle Cell"
row.cells[2].text = "Right Cell"

# Add dummy text to page
p = doc.add_paragraph("Hello world. This is a test of page borders and footer tables.")
doc.save("c:/Users/User/Desktop/antigravity/scratch/evolutio-game/scratch/test.docx")
print("Saved test.docx successfully!")
